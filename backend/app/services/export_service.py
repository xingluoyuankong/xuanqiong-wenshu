# 导出服务 - 支持 TXT 和 DOCX 格式导出小说
from __future__ import annotations

import io
import logging
from datetime import datetime

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from ..models import Chapter, ChapterOutline, NovelProject
from .novel_text_format import (
    EXPORT_HEADER_END_MARKER,
    EXPORT_HEADER_MARKER,
    build_chapter_title,
    encode_export_metadata,
)

logger = logging.getLogger(__name__)


class ExportService:
    """处理小说导出服务"""

    def __init__(self, session: AsyncSession):
        self.session = session

    async def export_novel_as_txt(self, project_id: str) -> str:
        """导出小说为 TXT 格式"""
        project = await self._get_project(project_id)
        chapters = await self._get_ordered_chapters(project_id)

        if not chapters:
            raise HTTPException(status_code=404, detail="没有章节可导出")
        self._validate_exportable_chapters(chapters)

        outlines = await self._get_outlines_map(project_id)

        output = []
        # 机读导出头标记，让导入侧可以确定性剥离书名/时间戳，不影响旧稿导入兼容性。
        output.append(EXPORT_HEADER_MARKER)
        output.append(encode_export_metadata(self._build_roundtrip_metadata(project, outlines, chapters=chapters)))
        output.append(f"书名: {project.title or '无标题小说'}")
        output.append(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        output.append(EXPORT_HEADER_END_MARKER)
        output.append("")

        for chapter in chapters:
            outline = outlines.get(chapter.chapter_number)
            chapter_title = build_chapter_title(
                chapter.chapter_number,
                outline.title if outline else None,
            )
            content = self._get_chapter_content(chapter)
            output.append(f"\n{chapter_title}")
            output.append(content)
            output.append("")

        return "\n".join(output)

    def _build_roundtrip_metadata(
        self,
        project: NovelProject,
        outlines: dict[int, ChapterOutline],
        *,
        chapters: list[Chapter] | None = None,
    ) -> dict:
        blueprint = getattr(project, "blueprint", None)
        blueprint_payload = None
        if blueprint is not None:
            world = dict(getattr(blueprint, "world_setting", None) or {})
            blueprint_payload = {
                "title": getattr(blueprint, "title", None) or project.title,
                "target_audience": getattr(blueprint, "target_audience", None) or "",
                "genre": getattr(blueprint, "genre", None) or "",
                "style": getattr(blueprint, "style", None) or "",
                "tone": getattr(blueprint, "tone", None) or "",
                "one_sentence_summary": getattr(blueprint, "one_sentence_summary", None) or "",
                "full_synopsis": getattr(blueprint, "full_synopsis", None) or "",
                "world_setting": world,
                "story_arcs": world.pop("story_arcs", []),
                "volume_plan": world.pop("volume_plan", []),
                "novel_outline": world.pop("novel_outline", []),
                "foreshadowing_system": world.pop("foreshadowing_system", []),
                "characters": [self._plain_character(item) for item in getattr(project, "characters", [])],
                "relationships": [self._plain_relationship(item) for item in getattr(project, "relationships_", [])],
            }
            blueprint_payload["world_setting"] = world
        formal_payload = {
            "foreshadowings": [self._plain_foreshadowing(item) for item in getattr(project, "foreshadowings", [])],
            "timeline_events": [self._plain_timeline_event(item) for item in getattr(project, "timeline_events", [])],
            "chapter_versions": [self._plain_chapter_versions(item) for item in (chapters or [])],
        }
        return {
            "format_version": 1,
            "project": {"title": getattr(project, "title", ""), "initial_prompt": getattr(project, "initial_prompt", ""), "status": getattr(project, "status", "")},
            "blueprint": blueprint_payload,
            "chapter_outlines": [self._plain_outline(outlines[n]) for n in sorted(outlines)],
            "formal_ledgers": formal_payload,
        }

    @staticmethod
    def _plain_character(item: object) -> dict:
        return {key: getattr(item, key, None) for key in ("name", "identity", "goals", "personality", "background", "arc", "description") if getattr(item, key, None) is not None}

    @staticmethod
    def _plain_relationship(item: object) -> dict:
        return {key: getattr(item, key, None) for key in ("character_from", "character_to", "description", "core_conflict", "relationship_type", "status", "tension", "direction", "trigger_event", "importance", "extra") if getattr(item, key, None) is not None}

    @staticmethod
    def _plain_outline(item: object) -> dict:
        keys = ("chapter_number", "title", "summary", "narrative_phase", "chapter_role", "suspense_hook", "emotional_progression", "character_focus", "cast_delta", "conflict_escalation", "continuity_notes", "foreshadowing", "foreshadowing_tasks", "payoff_window", "metadata")
        return {key: getattr(item, key, None) for key in keys if getattr(item, key, None) is not None}

    @staticmethod
    def _plain_foreshadowing(item: object) -> dict:
        keys = ("name", "content", "type", "keywords", "status", "chapter_number", "resolved_chapter_number",
                "target_reveal_chapter", "reveal_method", "reveal_impact", "related_characters", "related_plots",
                "importance", "urgency")
        return {key: getattr(item, key, None) for key in keys if getattr(item, key, None) is not None}

    @staticmethod
    def _plain_timeline_event(item: object) -> dict:
        keys = ("chapter_number", "story_time", "story_date", "time_elapsed", "event_type", "event_title",
                "event_description", "involved_characters", "location", "importance", "is_turning_point",
                "extra")
        return {key: getattr(item, key, None) for key in keys if getattr(item, key, None) is not None}

    @staticmethod
    def _plain_chapter_versions(item: object) -> dict:
        selected = getattr(item, "selected_version", None)
        versions = []
        for version in getattr(item, "versions", []) or []:
            versions.append({
                "version_label": getattr(version, "version_label", None),
                "provider": getattr(version, "provider", None),
                "content": getattr(version, "content", "") or "",
                "metadata": getattr(version, "metadata", None),
            })
        selected_index = None
        if selected is not None:
            for index, version in enumerate(getattr(item, "versions", []) or []):
                if getattr(version, "id", None) == getattr(selected, "id", None):
                    selected_index = index
                    break
        return {
            "chapter_number": getattr(item, "chapter_number", None),
            "versions": versions,
            "selected_index": selected_index,
        }


    async def preflight_export(self, project_id: str) -> dict:
        """导出前检查，给前端展示可执行的缺章/空章/未定稿原因。"""
        await self._get_project(project_id)
        chapters = await self._get_ordered_chapters(project_id)
        outlines = await self._get_outlines_map(project_id)
        issues = self._collect_export_issues(chapters)
        chapter_numbers = {int(getattr(chapter, "chapter_number", 0) or 0) for chapter in chapters}
        outline_numbers = {int(number) for number in outlines.keys()}
        missing_chapter_numbers = sorted(number for number in outline_numbers if number not in chapter_numbers)

        if not chapters:
            issues.append("没有章节可导出")
        for number in missing_chapter_numbers:
            issues.append(f"第{number}章只有大纲，尚未生成或定稿正文")

        exportable_chapters = 0
        total_word_count = 0
        for chapter in chapters:
            selected_version = getattr(chapter, "selected_version", None)
            content = (getattr(selected_version, "content", "") or "").strip() if selected_version else ""
            if getattr(chapter, "status", None) == "successful" and content:
                exportable_chapters += 1
                total_word_count += len("".join(content.split()))

        return {
            "ready": not issues,
            "total_chapters": len(chapters),
            "outline_chapters": len(outlines),
            "exportable_chapters": exportable_chapters,
            "total_word_count": total_word_count,
            "missing_chapter_numbers": missing_chapter_numbers,
            "issues": issues,
        }

    async def export_novel_as_docx(self, project_id: str) -> bytes:
        """导出小说为 DOCX 格式"""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="DOCX 导出功能未安装，请运行: pip install python-docx"
            )

        project = await self._get_project(project_id)
        chapters = await self._get_ordered_chapters(project_id)

        if not chapters:
            raise HTTPException(status_code=404, detail="没有章节可导出")
        self._validate_exportable_chapters(chapters)

        outlines = await self._get_outlines_map(project_id)

        doc = Document()
        doc.add_heading(project.title or "无标题小说", level=0)
        doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")

        for chapter in chapters:
            outline = outlines.get(chapter.chapter_number)
            chapter_title = build_chapter_title(
                chapter.chapter_number,
                outline.title if outline else None,
            )
            doc.add_page_break()
            doc.add_heading(chapter_title, level=1)

            content = self._get_chapter_content(chapter)
            if content:
                for para in content.split("\n\n"):
                    if not para.strip():
                        continue
                    paragraph = doc.add_paragraph(para.strip())
                    for run in paragraph.runs:
                        run.font.size = Pt(12)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    async def _get_project(self, project_id: str) -> NovelProject:
        """获取项目信息"""
        result = await self.session.execute(
            select(NovelProject)
            .where(NovelProject.id == project_id)
            .options(
                selectinload(NovelProject.blueprint),
                selectinload(NovelProject.characters),
                selectinload(NovelProject.relationships_),
                selectinload(NovelProject.foreshadowings),
                selectinload(NovelProject.timeline_events),
            )
        )
        project = result.scalar_one_or_none()
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        return project

    async def _get_ordered_chapters(self, project_id: str) -> list[Chapter]:
        """获取按顺序排列的章节列表"""
        result = await self.session.execute(
            select(Chapter)
            .where(Chapter.project_id == project_id)
            .options(
                selectinload(Chapter.selected_version),
                selectinload(Chapter.versions),
            )
            .order_by(Chapter.chapter_number)
        )
        return list(result.scalars().all())

    async def _get_outlines_map(self, project_id: str) -> dict[int, ChapterOutline]:
        result = await self.session.execute(
            select(ChapterOutline)
            .where(ChapterOutline.project_id == project_id)
            .order_by(ChapterOutline.chapter_number)
        )
        outlines = result.scalars().all()
        return {outline.chapter_number: outline for outline in outlines}

    def _collect_export_issues(self, chapters: list[Chapter]) -> list[str]:
        """导出前硬校验：禁止把空章节/未选中版本伪装成正常导出。

        历史实现会在 selected_version 缺失时回退到最新版本，导致数据库状态断链被
        TXT/DOCX 导出掩盖。小说交付场景中这会让作者误以为全书已完成，因此导出必须
        只接受“章节状态成功 + 已选中版本 + 正文非空”的章节。
        """
        invalid: list[str] = []
        for chapter in chapters:
            chapter_no = getattr(chapter, "chapter_number", "?")
            selected_version = getattr(chapter, "selected_version", None)
            content = (getattr(selected_version, "content", "") or "").strip() if selected_version else ""
            if getattr(chapter, "status", None) != "successful":
                invalid.append(f"第{chapter_no}章状态为 {getattr(chapter, 'status', 'unknown')}")
                continue
            if selected_version is None:
                invalid.append(f"第{chapter_no}章未选中正文版本")
                continue
            if not content:
                invalid.append(f"第{chapter_no}章选中版本正文为空")

        return invalid

    def _validate_exportable_chapters(self, chapters: list[Chapter]) -> None:
        invalid = self._collect_export_issues(chapters)
        if invalid:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "novel_export_not_ready",
                    "message": "小说仍存在未完成或不可导出的章节，请修复后再导出。",
                    "issues": invalid,
                },
            )

    def _get_chapter_content(self, chapter: Chapter) -> str:
        selected_version = chapter.selected_version
        if selected_version and selected_version.content:
            return selected_version.content

        return ""
